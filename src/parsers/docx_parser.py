from pathlib import Path

from docx import Document


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX including tables (pipe-delimited)."""
    doc = Document(Path(file_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)
